
from __future__ import annotations

from io import BytesIO
from typing import Optional

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt
from markdown_it import MarkdownIt
from markdown_it.token import Token
from mdit_py_plugins.gfm import gfm_plugin

_md = MarkdownIt("commonmark").use(gfm_plugin)

_LINK_COLOR = "2563EB"
_CODE_FONT = "Consolas"


def markdown_to_docx_bytes(markdown_text: str) -> bytes:
    document = Document()
    tokens = _md.parse(markdown_text)
    _render_tokens(document, tokens)
    buffer = BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def _render_tokens(document: Document, tokens: list[Token]) -> None:
    list_stack: list[str] = []  # one entry per nesting level: "bullet" | "ordered"
    seen_first_heading = False
    i = 0
    n = len(tokens)

    while i < n:
        tok = tokens[i]

        if tok.type == "heading_open":
            level = int(tok.tag[1:])
            inline = tokens[i + 1]
            if not seen_first_heading and level == 1:
                heading = document.add_heading(level=0)  # "Title" style
            else:
                heading = document.add_heading(level=min(level, 9))
            seen_first_heading = True
            _add_inline(heading, inline.children or [])
            i += 3
            continue

        if tok.type == "paragraph_open":
            inline = tokens[i + 1]
            if list_stack:
                paragraph = document.add_paragraph(style=_list_style(list_stack[-1], len(list_stack)))
            else:
                paragraph = document.add_paragraph()
            _add_inline(paragraph, inline.children or [])
            i += 3
            continue

        if tok.type == "bullet_list_open":
            list_stack.append("bullet")
            i += 1
            continue
        if tok.type == "ordered_list_open":
            list_stack.append("ordered")
            i += 1
            continue
        if tok.type in ("bullet_list_close", "ordered_list_close"):
            if list_stack:
                list_stack.pop()
            i += 1
            continue
        if tok.type in ("list_item_open", "list_item_close"):
            i += 1
            continue

        if tok.type == "blockquote_open":
            j = i + 1
            depth = 1
            while j < n and depth > 0:
                t = tokens[j]
                if t.type == "blockquote_open":
                    depth += 1
                elif t.type == "blockquote_close":
                    depth -= 1
                    if depth == 0:
                        break
                elif t.type == "inline":
                    paragraph = document.add_paragraph(style="Intense Quote")
                    _add_inline(paragraph, t.children or [])
                j += 1
            i = j + 1
            continue

        if tok.type == "fence" or tok.type == "code_block":
            _add_code_block(document, tok.content)
            i += 1
            continue

        if tok.type == "hr":
            _add_horizontal_rule(document.add_paragraph())
            i += 1
            continue

        if tok.type == "table_open":
            j, header_row, body_rows = _collect_table(tokens, i + 1)
            _add_table(document, header_row, body_rows)
            i = j + 1
            continue

        i += 1


def _list_style(kind: str, depth: int) -> str:
    base = "List Bullet" if kind == "bullet" else "List Number"
    if depth <= 1:
        return base
    return f"{base} {min(depth, 3)}"


def _add_inline(paragraph, children: list[Token], base_bold: bool = False, base_italic: bool = False) -> None:
    bold_depth = 0
    italic_depth = 0
    link_href: Optional[str] = None

    for tok in children:
        if tok.type == "strong_open":
            bold_depth += 1
        elif tok.type == "strong_close":
            bold_depth = max(0, bold_depth - 1)
        elif tok.type == "em_open":
            italic_depth += 1
        elif tok.type == "em_close":
            italic_depth = max(0, italic_depth - 1)
        elif tok.type == "link_open":
            link_href = tok.attrs.get("href")
        elif tok.type == "link_close":
            link_href = None
        elif tok.type == "code_inline":
            run = paragraph.add_run(tok.content)
            run.font.name = _CODE_FONT
            run.bold = base_bold or bold_depth > 0
            run.italic = base_italic or italic_depth > 0
        elif tok.type in ("softbreak", "hardbreak"):
            paragraph.add_run().add_break()
        elif tok.type in ("text", "html_inline") and tok.content:
            if link_href:
                _add_hyperlink(paragraph, tok.content, link_href)
            else:
                run = paragraph.add_run(tok.content)
                run.bold = base_bold or bold_depth > 0
                run.italic = base_italic or italic_depth > 0
        # image tokens are intentionally not handled -- long-form generated
        # articles from this app never contain them.


def _add_hyperlink(paragraph, text: str, url: str) -> None:
    part = paragraph.part
    r_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)

    run_el = OxmlElement("w:r")
    run_pr = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), _LINK_COLOR)
    run_pr.append(color)
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    run_pr.append(underline)
    run_el.append(run_pr)

    text_el = OxmlElement("w:t")
    text_el.set(qn("xml:space"), "preserve")
    text_el.text = text
    run_el.append(text_el)

    hyperlink.append(run_el)
    paragraph._p.append(hyperlink)


def _add_code_block(document: Document, content: str) -> None:
    paragraph = document.add_paragraph(style="No Spacing")
    paragraph.paragraph_format.left_indent = Pt(18)
    lines = content.rstrip("\n").split("\n")
    for idx, line in enumerate(lines):
        if idx > 0:
            paragraph.add_run().add_break()
        run = paragraph.add_run(line if line else " ")
        run.font.name = _CODE_FONT
        run.font.size = Pt(10)


def _add_horizontal_rule(paragraph) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "999999")
    p_bdr.append(bottom)
    p_pr.append(p_bdr)


def _collect_table(tokens: list[Token], start: int) -> tuple[int, Optional[list], list]:
    """Walk from just after table_open to table_close.

    Returns (index_of_table_close, header_row_cells_or_None, body_rows).
    Each "row" is a list of inline-children lists, one per cell.
    """
    header_row: Optional[list] = None
    body_rows: list = []
    current_row: list = []
    in_header = False
    j = start
    while j < len(tokens) and tokens[j].type != "table_close":
        t = tokens[j]
        if t.type == "thead_open":
            in_header = True
        elif t.type == "thead_close":
            in_header = False
        elif t.type == "tr_open":
            current_row = []
        elif t.type == "tr_close":
            if in_header:
                header_row = current_row
            else:
                body_rows.append(current_row)
        elif t.type == "inline":
            current_row.append(t.children or [])
        j += 1
    return j, header_row, body_rows


def _add_table(document: Document, header_row: Optional[list], body_rows: list) -> None:
    all_rows = ([header_row] if header_row else []) + body_rows
    if not all_rows:
        return
    num_cols = max(len(row) for row in all_rows)
    if num_cols == 0:
        return

    table = document.add_table(rows=0, cols=num_cols)
    table.style = "Table Grid"

    if header_row:
        cells = table.add_row().cells
        for idx, children in enumerate(header_row[:num_cols]):
            _add_inline(cells[idx].paragraphs[0], children, base_bold=True)

    for row in body_rows:
        cells = table.add_row().cells
        for idx, children in enumerate(row[:num_cols]):
            _add_inline(cells[idx].paragraphs[0], children)
