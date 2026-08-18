from io import BytesIO

from docx import Document

from backend.export import markdown_to_docx_bytes


def _load(markdown_text: str) -> Document:
    return Document(BytesIO(markdown_to_docx_bytes(markdown_text)))


def test_h1_becomes_title_style():
    doc = _load("# My Article\n\nSome text.")
    assert doc.paragraphs[0].style.name == "Title"
    assert doc.paragraphs[0].text == "My Article"


def test_h2_becomes_heading_2():
    doc = _load("# T\n\n## A Section\n\nBody.")
    headings = [p for p in doc.paragraphs if p.style.name == "Heading 2"]
    assert len(headings) == 1
    assert headings[0].text == "A Section"


def test_nested_bold_italic_runs_are_correct():
    doc = _load("Some **bold** and *italic* and ***both*** text.")
    runs = doc.paragraphs[0].runs
    by_text = {r.text: r for r in runs}
    assert by_text["bold"].bold is True
    assert by_text["bold"].italic is not True
    assert by_text["italic"].italic is True
    assert by_text["both"].bold is True
    assert by_text["both"].italic is True


def test_nested_bullet_list_uses_increasing_list_styles():
    doc = _load("- top level\n- also top\n  - nested one\n  - nested two")
    styles = [p.style.name for p in doc.paragraphs if p.text]
    assert styles == ["List Bullet", "List Bullet", "List Bullet 2", "List Bullet 2"]


def test_ordered_list_uses_list_number_style():
    doc = _load("1. first\n2. second")
    styles = [p.style.name for p in doc.paragraphs if p.text]
    assert styles == ["List Number", "List Number"]


def test_blockquote_uses_intense_quote_style():
    doc = _load("> a quoted line")
    assert doc.paragraphs[0].style.name == "Intense Quote"
    assert doc.paragraphs[0].text == "a quoted line"


def test_fenced_code_block_preserves_content_and_uses_monospace():
    doc = _load("```python\ndef f(x):\n    return x\n```")
    code_paragraph = next(p for p in doc.paragraphs if "def f(x):" in p.text)
    assert "return x" in code_paragraph.text
    assert code_paragraph.runs[0].font.name == "Consolas"


def test_link_produces_real_hyperlink_xml():
    doc = _load("Check out [this site](https://example.com) for more.")
    xml = doc.element.xml
    assert "w:hyperlink" in xml
    assert "https://example.com" in doc.part.rels[
        [r for r in doc.part.rels if doc.part.rels[r].reltype.endswith("hyperlink")][0]
    ].target_ref


def test_table_round_trips_with_header_and_all_rows():
    md = "| A | B |\n|---|---|\n| 1 | 2 |\n| 3 | 4 |\n"
    doc = _load(md)
    assert len(doc.tables) == 1
    table = doc.tables[0]
    assert [c.text for c in table.rows[0].cells] == ["A", "B"]
    assert [c.text for c in table.rows[1].cells] == ["1", "2"]
    assert [c.text for c in table.rows[2].cells] == ["3", "4"]


def test_horizontal_rule_adds_a_bordered_paragraph():
    doc = _load("above\n\n---\n\nbelow")
    assert "w:pBdr" in doc.element.xml


def test_full_sample_article_converts_without_error(sample_article_markdown):
    data = markdown_to_docx_bytes(sample_article_markdown)
    doc = Document(BytesIO(data))
    assert doc.paragraphs[0].style.name == "Title"
    assert len(doc.paragraphs) >= 8
